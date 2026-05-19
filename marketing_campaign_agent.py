"""
Marketing Campaign Agent
------------------------
1. Receives raw bytes of an attachment or OneDrive file plus filename/content-type.
2. Extracts readable text from PDF, DOCX, or plain-text files.
3. Asks GPT-4o with structured output whether it is a marketing product document.
4. If yes, returns a ready-to-use LinkedIn campaign message.
"""

import io
import json
import os
from typing import Optional

import openai
from pydantic import BaseModel

from config import OPENAI_API_KEY

openai.api_key = OPENAI_API_KEY

MARKETING_PROMPT_CONFIG_FILE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "marketing_campaign_system_prompt.json"
)

DEFAULT_MARKETING_SYSTEM_PROMPT = (
    "You are an expert B2B marketing campaign strategist. "
    "Read the provided extracted document text and decide whether it is "
    "a marketing product document. A marketing product document may be a "
    "product brief, product one-pager, launch document, brochure, sales "
    "enablement document, pitch deck text, feature announcement, case-study "
    "style product document, or similar document that describes a product, "
    "service, audience, benefits, features, positioning, or call to action."
)

MARKETING_EXTRACTION_INSTRUCTIONS = (
    "If the document is not product-marketing related, set "
    "is_marketing_product_document to false and all other fields to null. "
    "If it is product-marketing related, extract the product name, target "
    "audience, value proposition, key features, and campaign goal when present. "
    "Then write one polished LinkedIn campaign message. The message should be "
    "professional, clear, benefit-led, and ready to post or send as a LinkedIn "
    "campaign message. Do not invent hard metrics, prices, customer names, "
    "legal claims, or guarantees that are not supported by the document."
)


def get_marketing_system_prompt() -> str:
    """Return the saved custom prompt, or the built-in default if none is saved."""
    if os.path.exists(MARKETING_PROMPT_CONFIG_FILE):
        try:
            with open(MARKETING_PROMPT_CONFIG_FILE, encoding="utf-8") as f:
                data = json.load(f)
            text = data.get("system_prompt")
            if isinstance(text, str) and text.strip():
                return text
        except Exception as e:
            print(f"   Could not load marketing system prompt file: {e}")
    return DEFAULT_MARKETING_SYSTEM_PROMPT


def set_marketing_system_prompt(text: str) -> None:
    """Persist the marketing analysis system prompt (used by analyze_attachment)."""
    with open(MARKETING_PROMPT_CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump({"system_prompt": text}, f, ensure_ascii=False, indent=2)


class MarketingCampaignData(BaseModel):
    is_marketing_product_document: bool
    product_name: Optional[str]
    target_audience: Optional[str]
    value_proposition: Optional[str]
    key_features: Optional[str]
    campaign_goal: Optional[str]
    linkedin_message: Optional[str]


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception as e:
        print(f"   Could not read PDF: {e}")
        return ""


def _extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cleaned_text = cell.text.strip().replace("\n", " ")
                    if cleaned_text and cleaned_text not in row_data:
                        row_data.append(cleaned_text)
                if row_data:
                    full_text.append(" | ".join(row_data))

        return "\n".join(full_text)
    except Exception as e:
        print(f"   Could not read DOCX: {e}")
        return ""


def _extract_text_from_plain_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            text = text.replace("\x00", "").strip()
            if text:
                return text
        except UnicodeDecodeError:
            continue

    return file_bytes.decode("utf-8", errors="replace").replace("\x00", "").strip()


def extract_text(file_bytes: bytes, filename: str, content_type: str) -> str:
    """Route to the correct extractor based on file type."""
    name_lower = filename.lower()
    ct_lower = content_type.lower()
    text_extensions = (".txt", ".md", ".csv", ".json", ".rtf")

    if name_lower.endswith(".pdf") or "pdf" in ct_lower:
        return _extract_text_from_pdf(file_bytes)
    if (
        name_lower.endswith(".docx")
        or "wordprocessingml" in ct_lower
        or "msword" in ct_lower
    ):
        return _extract_text_from_docx(file_bytes)
    if (
        name_lower.endswith(text_extensions)
        or ct_lower.startswith("text/")
        or "json" in ct_lower
        or "csv" in ct_lower
    ):
        return _extract_text_from_plain_text(file_bytes)
    return ""


def analyze_attachment(
    file_bytes: bytes,
    filename: str,
    content_type: str,
    system_prompt: Optional[str] = None,
) -> Optional[MarketingCampaignData]:
    """
    Detects marketing product documents and generates a LinkedIn campaign message.
    Returns MarketingCampaignData when text can be analyzed, otherwise None.
    """
    text = extract_text(file_bytes, filename, content_type)
    print(f"   Extracted {len(text.strip()):,} characters of text from '{filename}'")
    if not text or len(text.strip()) < 50:
        print(f"   Attachment '{filename}' has no extractable text. Skipping.")
        return None

    truncated_text = text[:80000]
    print(
        f"   Sending {len(truncated_text):,} / {len(text):,} characters "
        "to AI for marketing analysis..."
    )

    effective_system_prompt = (
        system_prompt if system_prompt is not None else get_marketing_system_prompt()
    )

    # Ensure mandatory extraction instructions are present to maintain structured output quality
    if "is_marketing_product_document" not in effective_system_prompt:
        effective_system_prompt += "\n\n" + MARKETING_EXTRACTION_INSTRUCTIONS

    try:
        response = openai.beta.chat.completions.parse(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": effective_system_prompt,
                },
                {
                    "role": "user",
                    "content": f"Document filename: {filename}\n\nDocument text:\n{truncated_text}",
                },
            ],
            response_format=MarketingCampaignData,
            temperature=0.4,
        )
        return response.choices[0].message.parsed
    except Exception as e:
        print(f"   OpenAI error analyzing '{filename}': {e}")
        return None


def generate_campaign_from_text(
    description: str,
    system_prompt: Optional[str] = None,
) -> Optional[MarketingCampaignData]:
    """
    Generates a LinkedIn campaign message from a manual product description.
    """
    if not description or not description.strip():
        return None

    # Use a direct, forceful prompt for manual text entry
    instruction = """
    You are a marketing expert. Generate a professional LinkedIn marketing campaign based on the provided input.
    Even if the input is short (like a company name), use your knowledge to create a professional overview and a personalized DM.
    
    IMPORTANT RULES:
    1. Do NOT use placeholders like [Name], [Your Name], [Company], or [Date].
    2. Start the message naturally (e.g., "Hi there," or just start with a hook).
    3. Do NOT include a signature placeholder at the end.
    4. The message must be 100% ready to send as-is.
    
    Return the data in the specified JSON format.
    Always set is_marketing_product_document to true.
    """

    try:
        response = openai.beta.chat.completions.parse(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": instruction},
                {"role": "user", "content": f"Product Description: {description}"},
            ],
            response_format=MarketingCampaignData,
            temperature=0.7, # Slightly higher temperature for more creative generation from short names
        )
        return response.choices[0].message.parsed
    except Exception as e:
        print(f"❌ OpenAI Error in Manual Marketing: {e}")
        return None
